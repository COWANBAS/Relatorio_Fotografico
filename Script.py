import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

BASE_DIR = Path(__file__).resolve().parent
IMAGES_DIR = BASE_DIR / "Imagens"
DOCX_OUT = BASE_DIR / "Relatório Fotográfico.docx"

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', str(s))]

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def criar_cabecalho_formal(doc):
    col_width_total = Inches(7.5)
    w_logo = Inches(2.5)
    w_info = col_width_total - w_logo

    table = doc.add_table(rows=1, cols=2)
    table.style, table.autofit = 'Table Grid', False
    table.columns[0].width = w_logo
    table.columns[1].width = w_info

    c1 = table.cell(0, 0)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_background(c1, "1B2631")
    p_tit = c1.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_tit.add_run("RELATÓRIO\nFOTOGRÁFICO")
    run_tit.font.bold, run_tit.font.size, run_tit.font.name = True, Pt(14), 'Arial'
    run_tit.font.color.rgb = RGBColor(255, 255, 255) 
# Coloque entre as aspas cabeçalho.
    c2 = table.cell(0, 1)
    infos = [
        ("PROPONENTE:", ""),
        ("EMPREENDIMENTO:", ""),
        ("ENDEREÇO:", "")
    ]
    for i, (label, text) in enumerate(infos):
        p = c2.paragraphs[0] if i == 0 else c2.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r_lab = p.add_run(label)
        r_lab.font.bold, r_lab.font.size, r_lab.font.name = True, Pt(8), 'Arial'
        r_txt = p.add_run(f" {text}")
        r_txt.font.size, r_txt.font.name = Pt(8), 'Arial'

    tab_ax = doc.add_table(rows=1, cols=1)
    tab_ax.style = 'Table Grid'
    tab_ax.columns[0].width = col_width_total 
    cell_ax = tab_ax.cell(0, 0)
    set_cell_background(cell_ax, "F8F9F9")
    p_ax = cell_ax.paragraphs[0]
    p_ax.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ax = p_ax.add_run("REGISTRO FOTOGRÁFICO DE ACOMPANHAMENTO DE OBRA")
    run_ax.font.bold, run_ax.font.size, run_ax.font.name = True, Pt(11), 'Arial'

def inserir_foto(cell, path, largura, altura):
    """
    Insere a foto com dimensões fixas para garantir o alinhamento visual.
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p_img = cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(0)
    
    try:
        run_img = p_img.add_run()
        run_img.add_picture(str(path), width=largura, height=altura)
    except Exception:
        p_img.add_run(f"\n[ERRO NA IMAGEM: {path.name}]")

    p_legenda = cell.add_paragraph()
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_legenda.paragraph_format.space_before = Pt(6)
    p_legenda.paragraph_format.space_after = Pt(6)

    texto_legenda = path.stem.upper()
    run_leg = p_legenda.add_run(texto_legenda)
    run_leg.font.size = Pt(7.5)
    run_leg.font.bold = True
    run_leg.font.name = 'Arial'

def main():
    if not IMAGES_DIR.exists():
        print(f"ERRO: Pasta '{IMAGES_DIR.name}' não encontrada.")
        return

    exts = ('.png', '.jpg', '.jpeg')
    fotos = sorted([f for f in IMAGES_DIR.rglob('*') if f.suffix.lower() in exts], 
                   key=lambda x: natural_sort_key(x.name))
    if not fotos:
        print("Nenhuma imagem encontrada.")
        return
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.4)

    criar_cabecalho_formal(doc)
    doc.add_paragraph() 

    table_fotos = doc.add_table(rows=0, cols=2)
    table_fotos.style = 'Table Grid'
    table_fotos.autofit = False
    
    largura_coluna = Inches(3.7)
    table_fotos.columns[0].width = largura_coluna
    table_fotos.columns[1].width = largura_coluna

    LARGURA_FOTO = Inches(3.3)
    ALTURA_FOTO = Inches(2.4)

    for i in range(0, len(fotos), 2):
        row_cells = table_fotos.add_row().cells
        
        inserir_foto(row_cells[0], fotos[i], LARGURA_FOTO, ALTURA_FOTO)
        
        if i + 1 < len(fotos):
            inserir_foto(row_cells[1], fotos[i+1], LARGURA_FOTO, ALTURA_FOTO)

    try:
        doc.save(DOCX_OUT)
        print(f"Relatório gerado com {len(fotos)} fotos.")
    except PermissionError:
        print(f"ERRO: Feche o arquivo '{DOCX_OUT.name}' antes de executar.")

if __name__ == "__main__":
    main()
